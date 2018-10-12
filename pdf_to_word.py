import os
from configparser import ConfigParser
from concurrent.futures import ProcessPoolExecutor
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager, process_pdf
from io import StringIO
from io import open
from docx import Document

def read_from_pdf(file_path):
    print(file_path)
    with open(file_path,'rb') as file:
        print('我在read_from_pdf中')
        resource_manager = PDFResourceManager()
        return_str = StringIO()
        lap_params = LAParams()

        device = TextConverter(resource_manager,return_str,
                               laparams=lap_params)
        process_pdf(resource_manager,device,file)
        device.close()
        content = return_str.getvalue()
        return_str.close()
        return content
def save_text_to_word(content,file_path):
    print('我在save_text_to_word中')
    doc = Document()
    for line in content.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)

def remove_control_characters(content):
    mpa = dict.fromkeys(range(32))
    return content.translate(mpa)
def pdf_to_word(pdf_file_path,word_file_path):
    print('我在pdf_to_word中')
    content = read_from_pdf(pdf_file_path)
    print('我在pdf_to_word中2')
    save_text_to_word(content,word_file_path)
def main():
    config_parser = ConfigParser()
    config_parser.read('config.cfg')
    config = config_parser['default']
    tasks = []
    with ProcessPoolExecutor(max_workers=int(config['max_worker'])) as executor:
        for file in os.listdir(config['pdf_folder']):
            print('file:',file)
            extension_name = os.path.splitext(file)[1]
            if extension_name != '.pdf':
                continue
            file_name = os.path.splitext(file)[0]
            pdf_file = config['pdf_folder'] + '\\' +file
            print('pdf_file:',pdf_file)
            word_file = config['word_folder'] + '\\' + file_name + '.docx'
            print('word_file:',word_file)
            print('正在处理:',file)
            result = executor.submit(pdf_to_word,pdf_file,word_file)
            tasks.append(result)
        while True:
            exit_flag = True
            for task in tasks:
                if not task.done():
                    exit_flag = False
            if exit_flag:
                print('完成')
                exit(0)
if __name__ == '__main__':
    main()